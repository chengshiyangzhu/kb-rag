"""Shared pytest fixtures for the kb-rag test suite."""
from __future__ import annotations

from pathlib import Path
from typing import Any
from unittest.mock import MagicMock

import pytest


@pytest.fixture()
def tmp_chroma_path(tmp_path: Path) -> Path:
    """Return a temporary directory path for a Chroma store."""
    path = tmp_path / "chroma"
    path.mkdir(parents=True, exist_ok=True)
    return path


@pytest.fixture()
def mock_settings(tmp_chroma_path: Path) -> MagicMock:
    """Return a mock Settings object suitable for unit tests.

    Uses chroma as the vector store pointing at a temp dir so no network or
    persistent state is required.
    """
    settings = MagicMock()
    settings.app_env = "dev"
    settings.log_level = "INFO"
    settings.embedder_provider = "local"
    settings.embedder_model = "BAAI/bge-m3"
    settings.embedder_dim = 1024
    settings.ollama_base_url = "http://ollama:11434"
    settings.openai_api_key = "test-key"
    settings.openai_base_url = "https://api.openai.com/v1"
    settings.zhipu_api_key = "test-key"
    settings.zhipu_base_url = "https://open.bigmodel.cn/api/paas/v4"
    settings.llm_provider = "openai"
    settings.llm_model = "gpt-4o-mini"
    settings.vector_store = "chroma"
    settings.qdrant_url = "http://qdrant:6333"
    settings.qdrant_collection = "kb_rag_test"
    settings.chroma_path = str(tmp_chroma_path)
    settings.chunker_type = "recursive"
    settings.chunk_size = 512
    settings.chunk_overlap = 64
    settings.retrieve_top_n = 20
    settings.rerank_top_k = 5
    settings.rerank_model = "BAAI/bge-reranker-v2-m3"
    settings.rerank_threshold = 0.3
    settings.rrf_k = 60
    settings.bm25_index_path = str(tmp_chroma_path / "bm25.pkl")
    settings.prometheus_port = 9090
    settings.grafana_port = 3000
    settings.api_port = 8000
    settings.ui_port = 8501
    return settings


@pytest.fixture()
def sample_pdf(tmp_path: Path) -> Path:
    """Generate a small PDF file in tmp_path and return its path."""
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    path = tmp_path / "sample.pdf"
    with path.open("wb") as fh:
        writer.write(fh)
    return path


@pytest.fixture()
def sample_docx(tmp_path: Path) -> Path:
    """Generate a small DOCX file in tmp_path and return its path."""
    from docx import Document

    doc = Document()
    doc.add_heading("Sample Document", level=1)
    doc.add_paragraph("This is a sample paragraph for testing.")
    path = tmp_path / "sample.docx"
    doc.save(str(path))
    return path


@pytest.fixture()
def sample_xlsx(tmp_path: Path) -> Path:
    """Generate a small XLSX file in tmp_path and return its path."""
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws["A1"] = "Name"
    ws["B1"] = "Value"
    ws["A2"] = "alpha"
    ws["B2"] = 1
    ws["A3"] = "beta"
    ws["B3"] = 2
    path = tmp_path / "sample.xlsx"
    wb.save(str(path))
    return path


@pytest.fixture()
def sample_markdown(tmp_path: Path) -> Path:
    """Generate a small Markdown file in tmp_path and return its path."""
    path = tmp_path / "sample.md"
    path.write_text("# Sample\n\nThis is a **sample** markdown document.\n", encoding="utf-8")
    return path


@pytest.fixture()
def sample_files(sample_pdf: Path, sample_docx: Path, sample_xlsx: Path, sample_markdown: Path) -> dict[str, Path]:
    """Return a mapping of file-type to sample path for all supported types."""
    return {
        "pdf": sample_pdf,
        "docx": sample_docx,
        "xlsx": sample_xlsx,
        "md": sample_markdown,
    }


@pytest.fixture(autouse=True)
def _reset_settings_cache() -> Any:
    """Clear the lru_cache on app.config.get_settings between tests."""
    try:
        from app.config import get_settings
        get_settings.cache_clear()
    except Exception:
        # config import may fail if optional deps are missing; ignore in tests.
        pass
    yield
    try:
        from app.config import get_settings
        get_settings.cache_clear()
    except Exception:
        pass
